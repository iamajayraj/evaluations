import fitz  # PyMuPDF
import os
import pandas as pd
from io import BytesIO
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.embeddings import init_embeddings
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_core.documents import Document
from docx import Document as DocxDocument

from dotenv import load_dotenv
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

def table_name(page, table):
    finder = page.get_text("blocks")
    table_block = [block for block in finder if block[1]<table.bbox[1]]
    if table_block:
        return table_block[-1][4].strip()
    else:
        return None

def parse_pdf(file_bytes: bytes):

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    parsed_content = []
    for pno, page in enumerate(doc):

        # Extract full text
        page_text = page.get_text("text")

        #Extract tables
        page_tables = page.find_tables()
        tables = []
        for t_idx, table in enumerate(page_tables.tables):
            cells = table.extract()[:]
            # Add table name if available
            if table_name(page, table):
                table_title = table_name(page, table)
            else:
                table_title = "Table Title Unavailable"
            table_data = {"table_index": t_idx,
                            "table_title": table_title, 
                          "cells": cells,
                          "table": table.to_pandas().astype(str).to_markdown(index=False)}
            tables.append(table_data)

        # Extract images
        images = page.get_images(full=True)
        image_data = []
        for img_idx,img in enumerate(images):
            img_info = doc.extract_image(img[0])
            img_info = {"image_index": img_idx,
                        "image_bytes":img_info["image"],
                        "image_extension": img_info["ext"]}
            image_data.append(img_info)

        parsed_content.append({
            "page_number": pno,
            "page_text": page_text,
            "page_tables": tables,
            "page_images": image_data
        })

    return parsed_content

def parse_docx(file_bytes):
    doc = DocxDocument(BytesIO(file_bytes))
    return "\n\n".join(p.text for p in doc.paragraphs)


def bytes_to_markdown(file_bytes: bytes, suffix: str):
    "Convert excel and csv files to markdown tables."
    if suffix in [".xls", ".xlsx"]:
        df = pd.read_excel(BytesIO(file_bytes))
        return df.to_markdown(index=False, tablefmt="github")
    
    elif suffix == ".csv":
        df = pd.read_csv(BytesIO(file_bytes))
        return df.to_markdown(index=False, tablefmt="github")


class DocumentProcessor:
    def __init__(self):
        self._text_embeddings = None
        self._table_embeddings = None
        self.textual_retriever = None
        self.table_retriever = None
    
    def get_text_embeddings(self):
        if self._text_embeddings is None:
            self._text_embeddings = init_embeddings("openai:text-embedding-3-small", api_key=OPENAI_API_KEY)
        return self._text_embeddings
    
    def get_table_embeddings(self):
        if self._table_embeddings is None:
            self._table_embeddings = init_embeddings("openai:text-embedding-3-small", api_key=OPENAI_API_KEY)
        return self._table_embeddings
    
    def get_text_retriever(self, text: str):
        docs = [Document(page_content=text, metadata={"source": "inline"})]
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = text_splitter.split_documents(docs)
        embeddings = self.get_text_embeddings()
        vectorstore = InMemoryVectorStore.from_documents(documents=docs, embedding=embeddings)
        self.textual_retriever = vectorstore.as_retriever()
        return self.textual_retriever
    
    def get_table_retriever(self, tables_md: list[str]):
        docs = [Document(page_content=tbl, metadata={"chunk_type": "table", "table_index": i})
        for i, tbl in enumerate(tables_md)
        ]
        embeddings = self.get_table_embeddings()
        vectorstore = InMemoryVectorStore.from_documents(documents=docs, embedding=embeddings)
        self.table_retriever = vectorstore.as_retriever()
        return self.table_retriever


def get_text_retriever(text: str):
    processor = DocumentProcessor()
    return processor.get_text_retriever(text)

def get_table_retriever(tables_md: list[str]):
    processor = DocumentProcessor()
    return processor.get_table_retriever(tables_md)

def get_context(query, retriever):
    results = retriever.invoke(query)
    docs = "\n\n".join(doc.page_content for doc in results)
    return docs
